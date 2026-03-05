"""Tests for text normalisation, docx extraction, and chunking."""

import os
import tempfile
import pytest
from docx import Document
from backend.text_processing import (
    normalise_text,
    extract_docx_text,
    split_sentences,
    chunk_text,
    word_count,
)


# ---------------------------------------------------------------------------
# Normalisation tests
# ---------------------------------------------------------------------------

class TestNormaliseText:
    def test_replace_crlf(self):
        assert normalise_text("hello\r\nworld") == "hello\nworld"

    def test_replace_cr(self):
        assert normalise_text("hello\rworld") == "hello\nworld"

    def test_collapse_multiple_newlines(self):
        result = normalise_text("hello\n\n\n\nworld")
        assert result == "hello\n\nworld"

    def test_strip_line_whitespace(self):
        result = normalise_text("  hello  \n  world  ")
        assert result == "hello\nworld"

    def test_collapse_multiple_spaces(self):
        result = normalise_text("hello    world")
        assert result == "hello world"

    def test_collapse_tabs(self):
        result = normalise_text("hello\t\tworld")
        assert result == "hello world"

    def test_empty_string(self):
        assert normalise_text("") == ""

    def test_combined_normalisation(self):
        text = "  Hello  \r\n\r\n\r\n  World  \t here  "
        result = normalise_text(text)
        assert result == "Hello\n\nWorld here"


# ---------------------------------------------------------------------------
# Sentence splitting tests
# ---------------------------------------------------------------------------

class TestSplitSentences:
    def test_basic_split(self):
        result = split_sentences("Hello world. How are you? Fine!")
        assert result == ["Hello world.", "How are you?", "Fine!"]

    def test_single_sentence(self):
        result = split_sentences("Just one sentence.")
        assert result == ["Just one sentence."]

    def test_empty_string(self):
        result = split_sentences("")
        assert result == []

    def test_no_punctuation(self):
        result = split_sentences("No ending punctuation")
        assert result == ["No ending punctuation"]


# ---------------------------------------------------------------------------
# Chunking tests
# ---------------------------------------------------------------------------

class TestChunkText:
    def test_short_text_single_chunk(self):
        text = "This is a short text."
        chunks = chunk_text(text, chunk_size=100, overlap=10)
        assert len(chunks) == 1
        assert chunks[0]["text"] == "This is a short text."

    def test_empty_text(self):
        assert chunk_text("") == []
        assert chunk_text("   ") == []

    def test_chunks_have_required_keys(self):
        text = "First sentence here. Second sentence here. Third sentence here."
        chunks = chunk_text(text, chunk_size=5, overlap=1)
        for chunk in chunks:
            assert "text" in chunk
            assert "start" in chunk
            assert "end" in chunk
            assert isinstance(chunk["start"], int)
            assert isinstance(chunk["end"], int)

    def test_multiple_chunks(self):
        # Generate text with enough words to require multiple chunks
        sentences = [f"Sentence number {i} has several words in it." for i in range(20)]
        text = " ".join(sentences)
        chunks = chunk_text(text, chunk_size=20, overlap=5)
        assert len(chunks) > 1

    def test_overlap_exists(self):
        sentences = [f"Word{i}" + "." for i in range(50)]
        text = " ".join(sentences)
        chunks = chunk_text(text, chunk_size=10, overlap=3)
        if len(chunks) >= 2:
            # Words from end of first chunk should appear in start of second
            first_words = chunks[0]["text"].split()
            second_words = chunks[1]["text"].split()
            overlap_words = first_words[-3:]
            assert any(w in second_words for w in overlap_words)

    def test_long_sentence_gets_split(self):
        # A single sentence longer than chunk_size
        words = [f"word{i}" for i in range(100)]
        text = " ".join(words) + "."
        chunks = chunk_text(text, chunk_size=20, overlap=5)
        assert len(chunks) > 1

    def test_character_offsets(self):
        text = "First sentence. Second sentence. Third sentence."
        chunks = chunk_text(text, chunk_size=100, overlap=0)
        assert len(chunks) == 1
        assert chunks[0]["start"] >= 0
        assert chunks[0]["end"] <= len(text)


# ---------------------------------------------------------------------------
# DOCX extraction tests
# ---------------------------------------------------------------------------

class TestExtractDocx:
    def test_extract_from_file(self):
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Second paragraph")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = f.name

        try:
            text = extract_docx_text(path)
            assert "Hello World" in text
            assert "Second paragraph" in text
        finally:
            os.unlink(path)

    def test_extract_from_bytes(self):
        doc = Document()
        doc.add_paragraph("Bytes test")

        buf = tempfile.SpooledTemporaryFile()
        doc.save(buf)
        buf.seek(0)
        content = buf.read()

        text = extract_docx_text(content)
        assert "Bytes test" in text

    def test_skips_empty_paragraphs(self):
        doc = Document()
        doc.add_paragraph("Content")
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.add_paragraph("More content")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = f.name

        try:
            text = extract_docx_text(path)
            assert text == "Content\n\nMore content"
        finally:
            os.unlink(path)


# ---------------------------------------------------------------------------
# Word count tests
# ---------------------------------------------------------------------------

class TestWordCount:
    def test_basic_count(self):
        assert word_count("one two three") == 3

    def test_empty_string(self):
        assert word_count("") == 0

    def test_multiple_spaces(self):
        assert word_count("one  two  three") == 3
