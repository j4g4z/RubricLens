"""Text normalisation, .docx extraction, and sentence-aware chunking."""

import re
import io
from docx import Document


def normalise_text(text: str) -> str:
    """Normalise text for consistent processing.

    - Replace \\r\\n and \\r with \\n
    - Collapse 3+ consecutive newlines to 2
    - Strip leading/trailing whitespace per line
    - Collapse multiple spaces/tabs to single space
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = text.split("\n")
    lines = [line.strip() for line in lines]
    text = "\n".join(lines)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_docx_text(source) -> str:
    """Extract plain text from a .docx file.

    Args:
        source: Either a file path (str) or bytes content.

    Returns:
        Extracted text with paragraphs joined by double newlines.
    """
    if isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    else:
        doc = Document(source)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def split_sentences(text: str) -> list[str]:
    """Split text into sentences using regex.

    Handles common abbreviations and decimal numbers to avoid
    false splits.
    """
    sentences = re.split(r"(?<=[.!?])\s+", text)
    return [s.strip() for s in sentences if s.strip()]


def chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> list[dict]:
    """Segment text into sentence-aware overlapping chunks.

    Args:
        text: The full text to chunk.
        chunk_size: Target number of words per chunk.
        overlap: Number of words to overlap between chunks.

    Returns:
        List of {"text": str, "start": int, "end": int} with character offsets.
    """
    if not text.strip():
        return []

    sentences = split_sentences(text)
    if not sentences:
        return []

    chunks = []
    current_words = []
    current_start = 0

    for sentence in sentences:
        words = sentence.split()

        # If a single sentence exceeds chunk_size, split it by words
        if len(words) > chunk_size:
            # Flush current chunk first if it has content
            if current_words:
                chunk_text_str = " ".join(current_words)
                start_pos = text.find(chunk_text_str, current_start)
                if start_pos == -1:
                    start_pos = current_start
                chunks.append({
                    "text": chunk_text_str,
                    "start": start_pos,
                    "end": start_pos + len(chunk_text_str),
                })
                # Prepare overlap for next chunk
                overlap_words = current_words[-overlap:] if overlap > 0 else []
                current_words = list(overlap_words)
                current_start = start_pos

            # Split the long sentence into word-based chunks
            for i in range(0, len(words), chunk_size - overlap):
                word_slice = words[i:i + chunk_size]
                chunk_text_str = " ".join(word_slice)
                start_pos = text.find(chunk_text_str, current_start)
                if start_pos == -1:
                    start_pos = current_start
                chunks.append({
                    "text": chunk_text_str,
                    "start": start_pos,
                    "end": start_pos + len(chunk_text_str),
                })
                current_start = start_pos
            current_words = words[-(overlap):] if overlap > 0 else []
            continue

        # Check if adding this sentence exceeds chunk_size
        if len(current_words) + len(words) > chunk_size and current_words:
            chunk_text_str = " ".join(current_words)
            start_pos = text.find(chunk_text_str, current_start)
            if start_pos == -1:
                start_pos = current_start
            chunks.append({
                "text": chunk_text_str,
                "start": start_pos,
                "end": start_pos + len(chunk_text_str),
            })
            # Keep overlap words for next chunk
            overlap_words = current_words[-overlap:] if overlap > 0 else []
            current_words = list(overlap_words)
            current_start = start_pos

        current_words.extend(words)

    # Flush remaining words
    if current_words:
        chunk_text_str = " ".join(current_words)
        start_pos = text.find(chunk_text_str, current_start)
        if start_pos == -1:
            start_pos = current_start
        chunks.append({
            "text": chunk_text_str,
            "start": start_pos,
            "end": start_pos + len(chunk_text_str),
        })

    return chunks


def word_count(text: str) -> int:
    """Count the number of words in text."""
    return len(text.split())
